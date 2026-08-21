"""Test bản quyền và việc không phơi mã nguồn ra ngoài giao diện.

Hai chuyện khác nhau, dễ lẫn:

* **Bản quyền** — tên tác giả phải có mặt ở chân mọi trang, trong trang Hướng
  dẫn, và trong thuộc tính tệp báo cáo xuất ra.
* **Ẩn nguồn** — máy chủ không được phát ra thứ gì cho phép dựng lại mã nguồn
  hay hợp đồng API: không ``/openapi.json``, không liệt kê thư mục ``/static``,
  không trả vết ngăn xếp ra trình duyệt.
"""

from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook

from app import main as web
from app.core import audit, ban_quyen, report
from app.core.phien import CauHinh

# Gốc kho mã nguồn = thư mục APP (chứa app/, tests/, LICENSE...).
GOC_KHO = Path(__file__).resolve().parents[1]


@pytest.fixture
def khach():
    # raise_server_exceptions=False: cần đi qua đúng bộ xử lý ngoại lệ của app
    # để kiểm phần thân trả về, chứ không để TestClient ném lại lỗi.
    with TestClient(web.app, raise_server_exceptions=False) as c:
        c.cookies.set(web.TEN_COOKIE, web.TOKEN)
        yield c


@pytest.fixture
def bao_cao(tmp_path):
    kq = audit.KetQuaDoiSoat(goc=str(tmp_path))
    kq.tien_do = {ut: audit.TongTheoUuTien(0, 0) for ut in audit.MUC_UU_TIEN}
    tt = report.ThongTinVanBan.tu_cau_hinh(
        CauHinh(ten_dang_bo="Đảng bộ Thử nghiệm", ma_co_so="001")
    )
    return (
        report.xuat_bao_cao(kq, tt, tmp_path / "ra"),
        report.xuat_phu_luc(kq, tt, tmp_path / "ra"),
    )


class TestDongBanQuyen:
    def test_dung_dinh_dang(self):
        assert ban_quyen.dong_ban_quyen() == "© 2026 @anhduc97"

    def test_ngu_canh_mau_co_du_truong(self):
        for khoa in ("ten_ung_dung", "ten_hieu", "tac_gia", "ban_quyen", "giay_phep"):
            assert ban_quyen.NGU_CANH_BAN_QUYEN[khoa]


class TestBanQuyenTrenGiaoDien:
    def test_chan_moi_trang_deu_co(self, khach):
        for duong_dan in ("/buoc/0", "/huong_dan"):
            assert "© 2026 @anhduc97" in khach.get(duong_dan).text

    def test_co_dau_hieu_meo_may(self, khach):
        assert "/static/img/meo-may.svg" in khach.get("/buoc/0").text

    def test_tep_dau_hieu_ton_tai_va_khong_goi_ra_ngoai(self, khach):
        tra = khach.get("/static/img/meo-may.svg")
        assert tra.status_code == 200
        # App chạy ngoại tuyến. Dấu hiệu phải vẽ hoàn toàn bằng hình học, không
        # nhúng ảnh ngoài và không gọi phông chữ từ Internet. (Chuỗi
        # "http://www.w3.org/2000/svg" là tên vùng tên XML, không phải địa chỉ
        # tải về, nên không tính.)
        for cam in ("<image", "@import", "@font-face", "xlink:href=", "url(http"):
            assert cam not in tra.text, f"dấu hiệu còn phụ thuộc ngoài: {cam}"


class TestBanQuyenTrongTepXuatRa:
    def test_docx_ghi_tac_gia(self, bao_cao):
        tt = Document(bao_cao[0]).core_properties
        assert tt.author == "Mèo máy @anhduc97"
        assert "© 2026" in tt.comments

    def test_xlsx_ghi_tac_gia(self, bao_cao):
        tt = load_workbook(bao_cao[1]).properties
        assert tt.creator == "Mèo máy @anhduc97"

    def test_than_van_ban_khong_co_dong_quang_cao(self, bao_cao):
        """Báo cáo là văn bản hành chính, không được nhét tên phần mềm vào."""
        doc = Document(bao_cao[0])
        chu = "\n".join(p.text for p in doc.paragraphs)
        for cam in ("Mèo máy", "anhduc97", "©"):
            assert cam not in chu


class TestAnMaNguon:
    @pytest.mark.parametrize("duong_dan", ["/openapi.json", "/docs", "/redoc"])
    def test_khong_phoi_so_do_api(self, khach, duong_dan):
        assert khach.get(duong_dan).status_code == 404

    def test_khong_liet_ke_thu_muc_static(self, khach):
        assert khach.get("/static/").status_code == 404

    @pytest.mark.parametrize(
        "duong_dan",
        [
            "/static/../main.py",
            "/static/../core/rename.py",
            "/static/..%2fmain.py",
        ],
    )
    def test_khong_lay_duoc_ma_nguon_qua_static(self, khach, duong_dan):
        tra = khach.get(duong_dan)
        assert tra.status_code >= 400 or "def " not in tra.text

    def test_loi_ngoai_du_tinh_khong_lo_vet_ngan_xep(self, khach, monkeypatch):
        """Vết ngăn xếp lộ đường dẫn tuyệt đối và cấu trúc mô-đun của máy."""

        def no(*_a, **_k):
            raise RuntimeError("hong that")

        monkeypatch.setattr(web.chan_doan, "dung_goi", no)
        tra = khach.get("/api/chan_doan")
        assert tra.status_code == 500
        than = tra.json()["loi"]
        assert "RuntimeError: hong that" in than
        assert "Traceback" not in than
        assert "app/core" not in than and "app\\core" not in than


class TestTepGiayPhep:
    """Kho mã nguồn phải kèm giấy phép, nếu không thì đẩy lên GitHub là mất quyền."""

    def test_co_tep_license(self):
        tep = GOC_KHO / "LICENSE"
        assert tep.is_file(), "thiếu tệp LICENSE ở gốc kho mã nguồn"
        chu = tep.read_text(encoding="utf-8")
        assert "Mèo máy" in chu and "anhduc97" in chu

    def test_co_tep_notice(self):
        tep = GOC_KHO / "NOTICE.md"
        assert tep.is_file(), "thiếu tệp NOTICE.md"

    def test_gitignore_chan_du_lieu_ca_nhan(self):
        """Thư mục app/data chứa manifest có số CCCD thật — không được đẩy lên."""
        tep = GOC_KHO / ".gitignore"
        assert tep.is_file(), "thiếu .gitignore"
        chu = tep.read_text(encoding="utf-8")
        for mau in ("app/data/", "*.xlsx", "*.log"):
            assert mau in chu, f"gitignore thiếu {mau!r}"
