"""Test xuất báo cáo .docx (thể thức NĐ 30) và phụ lục .xlsx (bước 7)."""

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from app.core import audit, report
from app.core.mainbook import DongDangVien
from app.core.phien import CauHinh
from app.core.rename import KHO_CHO, dat_ten
from app.core.tree import duong_dan_tuong_doi

VP = "38.168.053.000.001"
QLKH = "38.168.053.000.002"
CHI_BO = {
    "Chi bộ Văn phòng": ("A", VP),
    "Chi bộ Phòng quản lý khoa học": ("B", QLKH),
}


def dv(ma_id, ten, cccd, unit=VP, chi_bo="Chi bộ Văn phòng"):
    from app.core.vietnamese import dung_folder_name, pascal_case

    return DongDangVien(
        id=ma_id, so_dong=int(ma_id[2:]), name=ten, name_convert=pascal_case(ten),
        folder_name=dung_folder_name(cccd, ten), unit_folder=unit,
        cccd_id=cccd, chi_bo_dang_sinh_hoat=chi_bo,
    )


@pytest.fixture
def ket_qua(tmp_path):
    so_cai = [
        dv("ID01", "NGUYỄN VĂN A", "012345678901"),
        dv("ID02", "TRẦN THỊ B", "012345678902"),
        dv("ID03", "LÊ VĂN C", "012345678903", QLKH, "Chi bộ Phòng quản lý khoa học"),
    ]
    for d in so_cai:
        thu_muc = tmp_path / duong_dan_tuong_doi(d)
        thu_muc.mkdir(parents=True)
        for ma in (1, 2):
            (thu_muc / dat_ten(ma, 1, ".pdf")).write_bytes(b"x")
    cho = tmp_path / duong_dan_tuong_doi(so_cai[0]) / KHO_CHO
    cho.mkdir()
    (cho / dat_ten(87, 1, ".docx")).write_bytes(b"x")
    return audit.doi_soat(tmp_path, so_cai, CHI_BO), so_cai


@pytest.fixture
def tt():
    """Thể thức điền sẵn từ bước 0, giống hệt đường đi thật của giao diện."""
    mac_dinh = report.ThongTinVanBan.tu_cau_hinh(
        CauHinh(
            ten_dang_bo="Đảng bộ Viện Nông nghiệp Thanh Hóa",
            ten_cap_tren="Đảng bộ UBND tỉnh Thanh Hóa",
            ma_co_so="053",
        )
    )
    mac_dinh.ngay = date(2026, 8, 20)
    mac_dinh.ho_ten_ky = "Nguyễn Văn Bí Thư"
    mac_dinh.so_hieu = "Số 12-BC/ĐU"
    return mac_dinh


def chu_trong(doc: Document) -> str:
    phan = [p.text for p in doc.paragraphs]
    for b in doc.tables:
        for hang in b.rows:
            phan += [o.text for o in hang.cells]
    return "\n".join(phan)


class TestBaoCaoDocx:
    def test_mo_duoc_va_du_chin_thanh_phan_the_thuc(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        tep = report.xuat_bao_cao(kq, tt, tmp_path / "ra")
        assert tep.is_file() and tep.suffix == ".docx"

        chu = chu_trong(Document(tep))
        assert "ĐẢNG CỘNG SẢN VIỆT NAM" in chu          # tiêu đề văn bản của Đảng
        assert "ĐẢNG ỦY VIỆN NÔNG NGHIỆP THANH HÓA" in chu
        assert "Số 12-BC/ĐU" in chu
        assert "Thanh Hóa, ngày 20 tháng 08 năm 2026" in chu
        assert "BÁO CÁO" in chu
        assert "Nơi nhận:" in chu
        assert "Lưu: VP Đảng ủy." in chu
        assert "T/M ĐẢNG ỦY" in chu
        assert "BÍ THƯ" in chu
        assert "Nguyễn Văn Bí Thư" in chu

    def test_du_bon_muc_lon_danh_so_nhat_quan(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        chu = chu_trong(Document(report.xuat_bao_cao(kq, tt, tmp_path / "ra")))
        for muc in ("I. CĂN CỨ", "II. KẾT QUẢ", "III. TỒN TẠI", "IV. NHIỆM VỤ"):
            assert muc in chu
        assert "1. Tổng quan" in chu and "2. Tiến độ theo mức độ ưu tiên" in chu

    def test_noi_ro_chua_dat_chuan_tt_02_2019(self, ket_qua, tt, tmp_path):
        """Không có câu này thì người đọc tưởng đã số hóa xong theo chuẩn lưu trữ."""
        kq, _ = ket_qua
        chu = chu_trong(Document(report.xuat_bao_cao(kq, tt, tmp_path / "ra")))
        assert "02/2019/TT-BNV" in chu
        assert "chưa được ký số" in chu
        assert "200 dpi" in chu

    def test_ket_thuc_bang_dau_gach_cheo(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        doc = Document(report.xuat_bao_cao(kq, tt, tmp_path / "ra"))
        cac_doan = [p.text for p in doc.paragraphs if p.text.strip()]
        assert cac_doan[-1].rstrip().endswith("./.")

    def test_so_lieu_trong_bang_khop_voi_doi_soat(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        chu = chu_trong(Document(report.xuat_bao_cao(kq, tt, tmp_path / "ra")))
        assert str(kq.tien_do[1].co) in chu
        assert "Chi bộ Văn phòng" in chu

    def test_the_thuc_nha_nuoc_thi_dung_quoc_hieu(self, ket_qua, tmp_path):
        kq, _ = ket_qua
        tt = report.ThongTinVanBan(the_thuc=report.THE_THUC_NHA_NUOC)
        chu = chu_trong(Document(report.xuat_bao_cao(kq, tt, tmp_path / "ra")))
        assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in chu
        assert "Độc lập - Tự do - Hạnh phúc" in chu

    def test_khong_dung_ky_tu_xuong_dong_trong_doan(self, ket_qua, tt, tmp_path):
        """Xuống dòng bằng \\n trong đoạn canh đều làm Word giãn chữ bất thường."""
        kq, _ = ket_qua
        doc = Document(report.xuat_bao_cao(kq, tt, tmp_path / "ra"))
        assert not any("\n" in p.text for p in doc.paragraphs)


class TestPhamViChiBo:
    def test_loc_con_mot_chi_bo(self, ket_qua):
        kq, _ = ket_qua
        con = report.loc_pham_vi(kq, QLKH)
        assert con.so_dang_vien == 1
        assert con.dong[0].id == "ID03"
        assert con.tien_do[1].tong == 36

    def test_bao_cao_mot_chi_bo_liet_ke_theo_dang_vien(self, ket_qua, tmp_path):
        kq, _ = ket_qua
        con = report.loc_pham_vi(kq, QLKH)
        tt = report.ThongTinVanBan(pham_vi=QLKH, ten_pham_vi="Chi bộ Phòng quản lý khoa học")
        chu = chu_trong(Document(report.xuat_bao_cao(con, tt, tmp_path / "ra")))
        assert "3. Kết quả theo từng đảng viên" in chu
        assert "LÊ VĂN C" in chu
        assert "NGUYỄN VĂN A" not in chu

    def test_ten_tep_mang_pham_vi_khong_dau(self, tmp_path):
        tt = report.ThongTinVanBan(ten_pham_vi="Chi bộ Văn phòng")
        ten = report.ten_tep_bao_cao(tt, ".docx")
        assert ten.startswith("BaoCao_SoHoaHSDV_Chi_bo_Van_phong_")
        assert ten.endswith(".docx")


class TestPhuLucXlsx:
    def test_du_ba_sheet(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        tep = report.xuat_phu_luc(kq, tt, tmp_path / "ra")
        wb = load_workbook(tep)
        assert wb.sheetnames == ["Tong_hop", "Chi_tiet_dang_vien", "Con_thieu_UT1"]

    def test_so_lieu_khop_voi_doi_soat(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        wb = load_workbook(report.xuat_phu_luc(kq, tt, tmp_path / "ra"))

        ws = wb["Tong_hop"]
        assert ws.cell(row=1, column=2).value == "Chi bộ"
        tong = [h for h in ws.iter_rows(values_only=True) if h[1] == "TỔNG CỘNG"][0]
        assert tong[3] == kq.so_dang_vien
        assert tong[4] == kq.tien_do[1].co

        ws2 = wb["Chi_tiet_dang_vien"]
        hang = list(ws2.iter_rows(min_row=2, values_only=True))
        assert len(hang) == 3
        assert hang[0][0] == "ID01"
        assert hang[0][7] == "1,2"
        assert hang[0][8] == "87"

    def test_sheet_con_thieu_chi_liet_ke_ut1(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        wb = load_workbook(report.xuat_phu_luc(kq, tt, tmp_path / "ra"))
        ws = wb["Con_thieu_UT1"]
        hang = list(ws.iter_rows(min_row=2, values_only=True))
        # Mỗi người có 1 và 2, còn thiếu 34 mã ƯT1 → 3 người × 34
        assert len(hang) == 3 * 34
        assert all(3 <= h[3] <= 36 for h in hang)
        assert all(h[4] for h in hang)      # có tên tiếng Việt của tài liệu

    def test_ten_tep_phu_luc(self, ket_qua, tt, tmp_path):
        kq, _ = ket_qua
        tep = report.xuat_phu_luc(kq, tt, tmp_path / "ra")
        assert tep.name.startswith("PhuLuc_SoHoaHSDV_")
        assert tep.suffix == ".xlsx"


class TestKhongKhoaCungDonVi:
    """Đặc tả DactaKetLuan: app phải dùng chung được cho đảng bộ khác.

    Cách hỏng dễ xảy ra nhất không phải là app chạy sai, mà là app chạy đúng và
    in ra tên đơn vị của người viết ra nó. Test này chạy với một đơn vị hư cấu
    rồi soi xem tên Viện Nông nghiệp có lọt vào tệp không.
    """

    @pytest.fixture
    def tt_don_vi_khac(self):
        mac_dinh = report.ThongTinVanBan.tu_cau_hinh(
            CauHinh(
                ten_dang_bo="Đảng bộ Trung tâm Khuyến nông Hà Tĩnh",
                ten_cap_tren="Đảng bộ UBND tỉnh Hà Tĩnh",
                ma_tinh="42",
                ma_cap_tren="170",
                ma_co_so="011",
                dia_danh="Hà Tĩnh",
            )
        )
        mac_dinh.ngay = date(2026, 8, 20)
        mac_dinh.ho_ten_ky = "Trần Văn B"
        return mac_dinh

    def test_docx_mang_ten_don_vi_moi(self, ket_qua, tt_don_vi_khac, tmp_path):
        kq, _ = ket_qua
        chu = chu_trong(
            Document(report.xuat_bao_cao(kq, tt_don_vi_khac, tmp_path / "ra"))
        )
        assert "ĐẢNG ỦY TRUNG TÂM KHUYẾN NÔNG HÀ TĨNH" in chu
        assert "ĐẢNG BỘ UBND TỈNH HÀ TĨNH" in chu
        assert "Hà Tĩnh, ngày 20 tháng 08 năm 2026" in chu

    def test_khong_con_dau_vet_vien_nong_nghiep(self, ket_qua, tt_don_vi_khac, tmp_path):
        kq, _ = ket_qua
        chu = chu_trong(
            Document(report.xuat_bao_cao(kq, tt_don_vi_khac, tmp_path / "ra"))
        )
        for dau_vet in ("Viện Nông nghiệp", "VIỆN NÔNG NGHIỆP", "Thanh Hóa"):
            assert dau_vet not in chu, f"còn sót {dau_vet!r} trong báo cáo"

    def test_can_cu_ke_hoach_100_ngay_theo_don_vi(self, ket_qua, tt_don_vi_khac, tmp_path):
        kq, _ = ket_qua
        chu = chu_trong(
            Document(report.xuat_bao_cao(kq, tt_don_vi_khac, tmp_path / "ra"))
        )
        assert "của Đảng bộ Trung tâm Khuyến nông Hà Tĩnh" in chu

    def test_thong_tin_van_ban_mac_dinh_khong_co_ten_don_vi_nao(self):
        """Mặc định phải rỗng, không được ghi sẵn tên đơn vị nào vào mã nguồn."""
        tt = report.ThongTinVanBan()
        assert tt.co_quan_ban_hanh == ""
        assert tt.co_quan_chu_quan == ""
        assert tt.ten_goi == ""
        assert tt.noi_nhan == []

    @pytest.mark.parametrize(
        "vao, ra",
        [
            ("Đảng bộ Viện Nông nghiệp Thanh Hóa", "ĐẢNG ỦY VIỆN NÔNG NGHIỆP THANH HÓA"),
            ("Chi bộ Trạm Khuyến nông X", "CHI BỘ TRẠM KHUYẾN NÔNG X"),
            ("", ""),
        ],
    )
    def test_doi_dang_bo_thanh_dang_uy(self, vao, ra):
        assert report.ten_co_quan_ban_hanh(vao) == ra
