"""Xuất báo cáo ``.docx`` theo Nghị định 30/2020/NĐ-CP và phụ lục ``.xlsx``.

Khối thể thức (Quốc hiệu — Tiêu ngữ, tên cơ quan có đường kẻ, địa danh — ngày
tháng, khối Nơi nhận — chữ ký) dựng lại theo bộ quy tắc trong kỹ năng
``soan-thao-van-ban-hanh-chinh`` của người dùng. Chép vào đây chứ không import
từ ngoài vì app phải chạy được độc lập, ngoại tuyến, trên máy trạm.

Hai điều đã học sẵn từ bộ quy tắc đó, giữ nguyên:

* Đường kẻ ngang dưới tên cơ quan nằm **trong một ô bảng**, nên phải truyền bề
  rộng ô thật (``container_width``); để mặc định theo bề rộng trang thì thụt lề
  vượt quá mép ô và đường kẻ co lại thành một chấm.
* Không dùng ``\\n`` trong đoạn văn canh đều — Word kéo giãn chữ ở dòng trước
  dấu xuống dòng. Mỗi ý là một đoạn riêng.

Báo cáo nói thẳng phần chưa đạt
-------------------------------
Mục "Tồn tại, hạn chế" luôn ghi rõ: v1 mới đặt tên và sắp xếp đúng quy tắc,
**chưa** ký số và chưa kiểm 200 dpi theo Thông tư 02/2019/TT-BNV. Không có mục
này thì người đọc dễ hiểu nhầm "đã số hóa xong theo chuẩn lưu trữ".
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from app.core import ban_quyen
from app.core.audit import MUC_UU_TIEN, KetQuaDoiSoat, TongTheoUuTien
from app.core.rename import danh_muc
from app.core.vietnamese import bo_dau

__all__ = [
    "THE_THUC_DANG",
    "THE_THUC_NHA_NUOC",
    "ThongTinVanBan",
    "ten_co_quan_ban_hanh",
    "loc_pham_vi",
    "ten_tep_bao_cao",
    "xuat_bao_cao",
    "xuat_phu_luc",
]

PHONG = "Times New Roman"

# Đơn vị ban hành là một tổ chức đảng nên mặc định dùng thể thức văn bản của
# Đảng: tiêu đề "ĐẢNG CỘNG SẢN VIỆT NAM" thay cho Quốc hiệu — Tiêu ngữ, số hiệu
# dạng "Số 12-BC/ĐU", người ký "T/M ĐẢNG ỦY". Bố cục, lề, phông và cách đánh đề
# mục vẫn theo NĐ 30/2020/NĐ-CP đúng như quyết định #10 của đặc tả. Ai cần bản
# thể thức nhà nước thì đổi một ô trên giao diện.
THE_THUC_DANG = "dang"
THE_THUC_NHA_NUOC = "nha_nuoc"

TEN_UU_TIEN = {
    1: "Ưu tiên 1 — tài liệu gốc, bắt buộc có trong hồ sơ",
    2: "Ưu tiên 2 — tài liệu quá trình công tác, khen thưởng, kỷ luật",
    3: "Ưu tiên 3 — tài liệu bổ trợ",
}


def ten_co_quan_ban_hanh(ten_dang_bo: str) -> str:
    """``Đảng bộ Viện Nông nghiệp`` ⇒ ``ĐẢNG ỦY VIỆN NÔNG NGHIỆP``.

    Tổ chức là *đảng bộ*, nhưng cơ quan ký ban hành văn bản là *đảng ủy* của
    đảng bộ đó. Đây chỉ là giá trị điền sẵn — ô này vẫn sửa được ở bước 7, vì
    một số đơn vị ban hành dưới tên chi bộ hoặc ban thường vụ.
    """
    ten = str(ten_dang_bo or "").strip()
    if not ten:
        return ""
    thap = ten.lower()
    if thap.startswith("đảng bộ "):
        ten = "Đảng ủy " + ten[len("Đảng bộ "):]
    return ten.upper()


@dataclass
class ThongTinVanBan:
    """Phần thể thức người dùng điền trên giao diện bước 7.

    Mặc định để **rỗng**, không ghi sẵn tên đơn vị nào. Tên và mã của đơn vị chỉ
    có một nguồn duy nhất là bước 0; viết thẳng tên một đơn vị vào đây là cách
    chắc chắn nhất để đơn vị thứ hai dùng app xuất ra báo cáo mang tên đơn vị
    thứ nhất mà không ai kịp nhận ra. Xem ``tu_cau_hinh``.
    """

    co_quan_chu_quan: str = ""
    co_quan_ban_hanh: str = ""
    so_hieu: str = "Số:        -BC/ĐU"
    dia_danh: str = ""
    ngay: date = field(default_factory=date.today)
    kinh_gui: str = ""
    quyen_han: str = "T/M ĐẢNG ỦY"
    chuc_vu_ky: str = "BÍ THƯ"
    ho_ten_ky: str = ""
    noi_nhan: list[str] = field(default_factory=list)
    don_vi_luu: str = "VP Đảng ủy"
    pham_vi: str = ""          # "" = toàn Đảng bộ, hoặc mã tổ chức của một chi bộ
    ten_pham_vi: str = "toàn Đảng bộ"
    ten_goi: str = ""
    the_thuc: str = THE_THUC_DANG
    can_cu_don_vi: str = ""    # tên đơn vị trong câu "Căn cứ Kế hoạch ... của ..."

    @classmethod
    def tu_cau_hinh(cls, ch) -> "ThongTinVanBan":
        """Dựng phần thể thức điền sẵn từ thông tin đơn vị khai ở bước 0."""
        ten_dang_bo = str(getattr(ch, "ten_dang_bo", "") or "").strip()
        ten_cap_tren = str(getattr(ch, "ten_cap_tren", "") or "").strip()
        noi_nhan = ["Các chi bộ trực thuộc"]
        if ten_cap_tren:
            noi_nhan.insert(0, f"{ten_cap_tren} (b/c)")
        return cls(
            co_quan_chu_quan=ten_cap_tren.upper(),
            co_quan_ban_hanh=ten_co_quan_ban_hanh(ten_dang_bo),
            dia_danh=str(getattr(ch, "dia_danh", "") or "").strip(),
            noi_nhan=noi_nhan,
            ten_goi=ten_dang_bo,
            can_cu_don_vi=ten_dang_bo,
        )


# ------------------------------------------------------- khối thể thức NĐ 30


def _dat_run(run, co=13, dam=False, nghieng=False, gach_chan=False):
    run.font.name = PHONG
    run.font.size = Pt(co)
    run.font.bold = dam
    run.font.italic = nghieng
    run.font.underline = gach_chan
    return run


def _ghi_tac_gia(thuoc_tinh) -> None:
    """Đặt tác giả vào thuộc tính của tệp Office (cả .docx lẫn .xlsx).

    Dấu vết bản quyền đi theo tệp mà không làm bẩn nội dung: phần thân báo cáo
    là văn bản hành chính của đảng ủy, tuyệt đối không được có dòng nào quảng
    cáo phần mềm. Ai cần biết tệp do đâu ra thì mở Thuộc tính của tệp.

    python-docx và openpyxl đặt tên trường khác nhau (``author`` với
    ``creator``), nên phải dò chứ không gán bừa — gán sai tên thì openpyxl im
    lặng tạo thuộc tính rác và tệp vẫn ghi "openpyxl" là người tạo.
    """
    ten = f"{ban_quyen.TEN_HIEU} {ban_quyen.TAC_GIA}"
    mo_ta = f"{ban_quyen.TEN_UNG_DUNG} — {ban_quyen.dong_ban_quyen()}"
    for truong, gia_tri in (
        ("author", ten),            # python-docx
        ("creator", ten),           # openpyxl
        ("last_modified_by", ten),  # python-docx
        ("lastModifiedBy", ten),    # openpyxl
        ("comments", mo_ta),        # python-docx
        ("description", mo_ta),     # openpyxl
    ):
        if hasattr(thuoc_tinh, truong):
            setattr(thuoc_tinh, truong, gia_tri)


def _dung_tai_lieu() -> Document:
    doc = Document()
    _ghi_tac_gia(doc.core_properties)
    khu = doc.sections[0]
    khu.top_margin = Mm(22)
    khu.bottom_margin = Mm(22)
    khu.left_margin = Mm(32)
    khu.right_margin = Mm(18)

    thuong = doc.styles["Normal"]
    thuong.font.name = PHONG
    thuong.font.size = Pt(13)
    rpr = thuong.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), PHONG)
    thuong.paragraph_format.line_spacing = 1.5
    thuong.paragraph_format.space_after = Pt(6)
    return doc


def _doan(doc, chu="", co=13, dam=False, nghieng=False,
          canh=WD_ALIGN_PARAGRAPH.JUSTIFY, cach_sau=6):
    p = doc.add_paragraph()
    p.alignment = canh
    p.paragraph_format.space_after = Pt(cach_sau)
    if chu:
        _dat_run(p.add_run(chu), co=co, dam=dam, nghieng=nghieng)
    return p


def _duong_ke(p, ty_le=0.4, be_rong_o=None):
    """Đường kẻ ngang căn giữa, ngắn hơn vùng chứa.

    ``be_rong_o`` bắt buộc khi đoạn nằm trong ô bảng — xem ghi chú đầu tệp.
    """
    if be_rong_o is None:
        khu = p.part.document.sections[0]
        be_rong_o = khu.page_width - khu.left_margin - khu.right_margin
    rong_ke = int(be_rong_o * ty_le)
    thut = (be_rong_o - rong_ke) // 2

    dinh_dang = p.paragraph_format
    dinh_dang.left_indent = thut
    dinh_dang.right_indent = thut
    dinh_dang.space_before = Pt(0)
    dinh_dang.space_after = Pt(2)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    duoi = OxmlElement("w:bottom")
    duoi.set(qn("w:val"), "single")
    duoi.set(qn("w:sz"), "6")
    duoi.set(qn("w:space"), "1")
    duoi.set(qn("w:color"), "auto")
    pBdr.append(duoi)
    pPr.append(pBdr)


def _be_rong_noi_dung(doc) -> int:
    khu = doc.sections[0]
    return khu.page_width - khu.left_margin - khu.right_margin


def _dat_be_rong_cot(bang, rong_cot: list[int]) -> None:
    """Ép bề rộng từng cột của bảng.

    Chỉ gán ``cell.width`` là chưa đủ: Word/LibreOffice vẫn tự co giãn theo nội
    dung nếu bảng còn ở chế độ autofit. Phải đặt cả ``tblLayout`` sang "fixed"
    và ghi bề rộng vào ``tblGrid`` — đây mới là chỗ LibreOffice đọc.
    """
    bang.autofit = False
    tbl = bang._tbl
    tblPr = tbl.tblPr
    bo_cuc = OxmlElement("w:tblLayout")
    bo_cuc.set(qn("w:type"), "fixed")
    tblPr.append(bo_cuc)

    luoi = tbl.find(qn("w:tblGrid"))
    if luoi is not None:
        for cot, rong in zip(luoi.findall(qn("w:gridCol")), rong_cot):
            cot.set(qn("w:w"), str(int(rong / 635)))     # EMU -> twip

    for hang in bang.rows:
        for o, rong in zip(hang.cells, rong_cot):
            o.width = rong


def _khoi_dau(doc, tt: ThongTinVanBan) -> None:
    bang = doc.add_table(rows=1, cols=2)
    bang.alignment = WD_TABLE_ALIGNMENT.CENTER
    bang.autofit = False
    trai, phai = bang.rows[0].cells

    # Cột phải rộng hơn: dòng "Thanh Hóa, ngày 20 tháng 08 năm 2026" chia đôi
    # đều nhau sẽ bị ngắt xuống dòng, nhìn rất nghiệp dư.
    tong_rong = _be_rong_noi_dung(doc)
    rong_trai = int(tong_rong * 0.46)
    _dat_be_rong_cot(bang, [rong_trai, tong_rong - rong_trai])
    be_rong_o = rong_trai

    if tt.co_quan_chu_quan:
        p = trai.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _dat_run(p.add_run(tt.co_quan_chu_quan.upper()), co=13)
        p_ten = trai.add_paragraph()
    else:
        p_ten = trai.paragraphs[0]
    p_ten.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ten.paragraph_format.space_after = Pt(0)
    _dat_run(p_ten.add_run(tt.co_quan_ban_hanh.upper()), co=13, dam=True)

    p_ke = trai.add_paragraph()
    p_ke.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _duong_ke(p_ke, ty_le=0.4, be_rong_o=be_rong_o)

    p_so = trai.add_paragraph()
    p_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dat_run(p_so.add_run(tt.so_hieu), co=13)

    p1 = phai.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    if tt.the_thuc == THE_THUC_DANG:
        # Văn bản của tổ chức đảng dùng tiêu đề "ĐẢNG CỘNG SẢN VIỆT NAM" thay
        # cho Quốc hiệu — Tiêu ngữ. Bố cục, lề, phông, cách đánh đề mục vẫn
        # theo NĐ 30 như đặc tả đã chốt.
        _dat_run(p1.add_run("ĐẢNG CỘNG SẢN VIỆT NAM"), co=14, dam=True)
        p2 = phai.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _duong_ke(p2, ty_le=0.55, be_rong_o=be_rong_o)
    else:
        _dat_run(p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"), co=13, dam=True)
        p2 = phai.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        _dat_run(
            p2.add_run("Độc lập - Tự do - Hạnh phúc"), co=14, dam=True, gach_chan=True
        )

    p_ngay = phai.add_paragraph()
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ngay.paragraph_format.space_before = Pt(6)
    _dat_run(
        p_ngay.add_run(
            f"{tt.dia_danh}, ngày {tt.ngay.day:02d} tháng {tt.ngay.month:02d} "
            f"năm {tt.ngay.year}"
        ),
        co=14,
        nghieng=True,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _khoi_ten_loai(doc, trich_yeu: str) -> None:
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    _dat_run(p1.add_run("BÁO CÁO"), co=14, dam=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dat_run(p2.add_run(trich_yeu), co=14, dam=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _duong_ke(p3, ty_le=0.30)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _khoi_ky(doc, tt: ThongTinVanBan) -> None:
    bang = doc.add_table(rows=1, cols=2)
    hang = bang.rows[0]
    trai, phai = hang.cells

    # Khóa dòng bảng không cho tách qua trang: nếu không, khối Nơi nhận — chữ ký
    # rơi vào cuối trang sẽ bị ngắt ngay giữa, rất mất mỹ quan.
    trPr = hang._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))

    p_nhan = trai.paragraphs[0]
    p_nhan.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _dat_run(p_nhan.add_run("Nơi nhận:"), co=12, dam=True, nghieng=True)

    cac_dong = list(tt.noi_nhan) + [f"Lưu: {tt.don_vi_luu}."]
    for i, dong in enumerate(cac_dong):
        p = trai.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        duoi = "" if dong.endswith((".", ";")) else (";" if i < len(cac_dong) - 1 else ".")
        _dat_run(p.add_run(f"- {dong}{duoi}"), co=11)

    p1 = phai.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    if tt.quyen_han:
        # "T/M ĐẢNG ỦY" và "BÍ THƯ" nằm hai dòng, không viết liền một dòng.
        _dat_run(p1.add_run(tt.quyen_han.upper()), co=13, dam=True)
        p_chuc = phai.add_paragraph()
        p_chuc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _dat_run(p_chuc.add_run(tt.chuc_vu_ky.upper()), co=13, dam=True)
    else:
        _dat_run(p1.add_run(tt.chuc_vu_ky.upper()), co=13, dam=True)
    for _ in range(3):
        phai.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = phai.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dat_run(p2.add_run(tt.ho_ten_ky or "…………………………"), co=13, dam=True)


def _bang(doc, tieu_de: list[str], cac_hang: list[list[str]],
          ty_le_cot: list[float] | None = None) -> None:
    """Bảng số liệu. ``ty_le_cot`` là phần trăm bề rộng từng cột, cộng lại bằng 1.

    Phải đặt bề rộng rõ ràng: để Word tự co, cột tên tài liệu dài bị bóp lại
    còn vài chữ trong khi cột số thì thừa chỗ.
    """
    bang = doc.add_table(rows=1, cols=len(tieu_de))
    bang.style = "Table Grid"
    bang.alignment = WD_TABLE_ALIGNMENT.CENTER
    rong_cot = (
        [int(_be_rong_noi_dung(doc) * t) for t in ty_le_cot] if ty_le_cot else []
    )
    for o, chu in zip(bang.rows[0].cells, tieu_de):
        p = o.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        _dat_run(p.add_run(chu), co=12, dam=True)
    for hang in cac_hang:
        cac_o = bang.add_row().cells
        for i, (o, chu) in enumerate(zip(cac_o, hang)):
            p = o.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            _dat_run(p.add_run(str(chu)), co=12)
    if rong_cot:
        _dat_be_rong_cot(bang, rong_cot)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ------------------------------------------------------------------ nội dung


def loc_pham_vi(kq: KetQuaDoiSoat, ma_to_chuc: str) -> KetQuaDoiSoat:
    """Cắt kết quả đối soát xuống còn một chi bộ. Trả về bản mới, không sửa bản gốc."""
    if not ma_to_chuc:
        return kq
    con = KetQuaDoiSoat(goc=kq.goc)
    con.dong = [d for d in kq.dong if d.unit_folder == ma_to_chuc]
    con.chi_bo = [c for c in kq.chi_bo if c.ma_to_chuc == ma_to_chuc]
    con.tien_do = {
        ut: TongTheoUuTien(
            sum(d.tien_do[ut].co for d in con.dong),
            sum(d.tien_do[ut].tong for d in con.dong),
        )
        for ut in MUC_UU_TIEN
    }
    con.so_cho_chuyen_pdf = sum(len(d.cho_chuyen_pdf) for d in con.dong)
    con.thieu_thu_muc = sum(0 if d.co_thu_muc else 1 for d in con.dong)
    return con


def ten_tep_bao_cao(tt: ThongTinVanBan, duoi: str) -> str:
    pham_vi = bo_dau(tt.ten_pham_vi).replace(" ", "_")
    return f"BaoCao_SoHoaHSDV_{pham_vi}_{datetime.now():%Y%m%d_%H%M%S}{duoi}"


def xuat_bao_cao(kq: KetQuaDoiSoat, tt: ThongTinVanBan, thu_muc_ra: Path) -> Path:
    """Dựng tệp .docx theo thể thức NĐ 30 và ghi ra đĩa."""
    doc = _dung_tai_lieu()
    _khoi_dau(doc, tt)
    _khoi_ten_loai(
        doc,
        f"Kết quả số hóa hồ sơ đảng viên {tt.ten_pham_vi} "
        f"(tính đến ngày {tt.ngay.day:02d}/{tt.ngay.month:02d}/{tt.ngay.year})",
    )
    if tt.kinh_gui:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _dat_run(p.add_run(f"Kính gửi: {tt.kinh_gui}"), co=13, dam=True)

    tong_ut = kq.tien_do

    _doan(doc, "I. CĂN CỨ VÀ PHẠM VI BÁO CÁO", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
    cua_don_vi = f" của {tt.can_cu_don_vi}" if tt.can_cu_don_vi else ""
    for cau in (
        f"Căn cứ Kế hoạch 100 ngày số hóa hồ sơ, tài liệu{cua_don_vi};",
        "Căn cứ Kế hoạch số hóa hồ sơ đảng viên và Phụ lục 1 kèm theo (danh mục "
        "104 loại tài liệu, chia ba mức độ ưu tiên);",
        "Căn cứ Thông tư số 02/2019/TT-BNV ngày 24/01/2019 của Bộ Nội vụ quy định "
        "tiêu chuẩn dữ liệu thông tin đầu vào và yêu cầu bảo quản tài liệu lưu trữ "
        "điện tử;",
    ):
        _doan(doc, cau, nghieng=True)
    _doan(
        doc,
        f"Phạm vi báo cáo: {tt.ten_pham_vi}, gồm {kq.so_dang_vien} đảng viên "
        f"thuộc {len(kq.chi_bo)} chi bộ.",
    )

    _doan(doc, "II. KẾT QUẢ THỰC HIỆN", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
    _doan(doc, "1. Tổng quan", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
    _doan(
        doc,
        f"Đã lập đủ hệ thống thư mục hồ sơ cho {kq.so_dang_vien} đảng viên theo "
        f"mã tổ chức đảng quy định tại Quy định số 208-QĐ/TW. Tổng số lượt tài "
        f"liệu đã số hóa và lưu đúng vị trí là {kq.tong_tep_da_co} tài liệu; "
        f"còn {kq.so_cho_chuyen_pdf} tài liệu đã sắp xếp đúng quy tắc nhưng chưa "
        f"ở định dạng PDF nên chưa được tính vào kết quả đạt chuẩn.",
    )

    _doan(doc, "2. Tiến độ theo mức độ ưu tiên", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
    _bang(
        doc,
        ["Mức độ", "Nhóm tài liệu", "Đã đạt chuẩn", "Tổng số cần có", "Tỷ lệ"],
        [
            [
                f"ƯT{ut}",
                TEN_UU_TIEN[ut],
                tong_ut[ut].co,
                tong_ut[ut].tong,
                f"{tong_ut[ut].ti_le:.1f}%",
            ]
            for ut in MUC_UU_TIEN
        ],
        ty_le_cot=[0.11, 0.45, 0.15, 0.16, 0.13],
    )

    if len(kq.chi_bo) > 1:
        _doan(doc, "3. Kết quả theo từng chi bộ", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
        _bang(
            doc,
            ["TT", "Chi bộ", "Số ĐV", "ƯT1 đạt", "Tỷ lệ ƯT1", "Chờ chuyển PDF"],
            [
                [
                    i,
                    c.ten,
                    c.so_dang_vien,
                    f"{c.tien_do[1].co}/{c.tien_do[1].tong}",
                    f"{c.tien_do[1].ti_le:.1f}%",
                    c.so_cho_chuyen_pdf,
                ]
                for i, c in enumerate(kq.chi_bo, 1)
            ],
            ty_le_cot=[0.07, 0.40, 0.10, 0.13, 0.14, 0.16],
        )
    else:
        _doan(doc, "3. Kết quả theo từng đảng viên", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
        _bang(
            doc,
            ["TT", "Họ và tên", "ƯT1", "ƯT2", "ƯT3", "Chờ chuyển PDF"],
            [
                [
                    i,
                    d.ho_ten,
                    str(d.tien_do[1]),
                    str(d.tien_do[2]),
                    str(d.tien_do[3]),
                    len(d.cho_chuyen_pdf),
                ]
                for i, d in enumerate(kq.dong, 1)
            ],
            ty_le_cot=[0.07, 0.37, 0.12, 0.12, 0.12, 0.20],
        )

    _doan(doc, "III. TỒN TẠI, HẠN CHẾ", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
    _doan(
        doc,
        "1. Kết quả nêu trên mới phản ánh việc thu thập, đặt tên và sắp xếp tài "
        "liệu theo đúng danh mục. Tài liệu chưa được ký số và chưa kiểm tra thông "
        "số quét (độ phân giải tối thiểu 200 dpi, ảnh màu) theo Thông tư số "
        "02/2019/TT-BNV, do đó chưa đủ điều kiện coi là tài liệu lưu trữ điện tử "
        "hoàn chỉnh.",
    )
    _doan(
        doc,
        f"2. Còn {kq.so_cho_chuyen_pdf} tài liệu ở định dạng Word hoặc ảnh, đang "
        f"được lưu riêng tại thư mục chờ chuyển đổi. Các tài liệu này phải được "
        f"chuyển sang PDF trước khi tính vào kết quả số hóa.",
    )
    if kq.thieu_thu_muc:
        _doan(
            doc,
            f"3. Còn {kq.thieu_thu_muc} đảng viên chưa có thư mục hồ sơ trên hệ "
            f"thống, cần rà soát lại danh sách và mã tổ chức đảng của chi bộ.",
        )

    _doan(doc, "IV. NHIỆM VỤ THỜI GIAN TỚI", dam=True, canh=WD_ALIGN_PARAGRAPH.LEFT)
    _doan(
        doc,
        f"1. Tập trung hoàn thành nhóm tài liệu ưu tiên 1, hiện đạt "
        f"{tong_ut[1].ti_le:.1f}%; giao các chi bộ rà soát, bổ sung tài liệu còn "
        f"thiếu theo phụ lục kèm theo báo cáo này.",
    )
    _doan(doc, "2. Chuyển toàn bộ tài liệu đang chờ sang định dạng PDF.")
    _doan(
        doc,
        "3. Xây dựng phương án ký số và kiểm tra thông số quét để bảo đảm yêu cầu "
        "của Thông tư số 02/2019/TT-BNV.",
    )
    _doan(
        doc,
        f"{tt.ten_goi} báo cáo kết quả số hóa hồ sơ đảng viên "
        f"{tt.ten_pham_vi} để cấp trên và các chi bộ biết, chỉ đạo thực hiện./.",
    )

    _khoi_ky(doc, tt)

    thu_muc_ra = Path(thu_muc_ra)
    thu_muc_ra.mkdir(parents=True, exist_ok=True)
    tep = thu_muc_ra / ten_tep_bao_cao(tt, ".docx")
    doc.save(tep)
    return tep


# ------------------------------------------------------------------- phụ lục


def _dat_tieu_de(ws, tieu_de: list[str], be_rong: list[int]) -> None:
    ws.append(tieu_de)
    for o in ws[1]:
        o.font = Font(bold=True)
        o.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for i, rong in enumerate(be_rong, 1):
        ws.column_dimensions[get_column_letter(i)].width = rong
    ws.freeze_panes = "A2"


def xuat_phu_luc(kq: KetQuaDoiSoat, tt: ThongTinVanBan, thu_muc_ra: Path) -> Path:
    """Phụ lục số liệu ba sheet, số liệu lấy đúng từ kết quả đối soát."""
    wb = Workbook()
    _ghi_tac_gia(wb.properties)

    ws = wb.active
    ws.title = "Tong_hop"
    _dat_tieu_de(
        ws,
        ["TT", "Chi bộ", "Mã tổ chức đảng", "Số đảng viên",
         "ƯT1 đạt", "ƯT1 tổng", "Tỷ lệ ƯT1 (%)",
         "ƯT2 đạt", "ƯT3 đạt", "Chờ chuyển PDF", "Thiếu thư mục"],
        [5, 34, 20, 12, 10, 10, 14, 10, 10, 16, 14],
    )
    for i, c in enumerate(kq.chi_bo, 1):
        ws.append([
            i, c.ten, c.ma_to_chuc, c.so_dang_vien,
            c.tien_do[1].co, c.tien_do[1].tong, round(c.tien_do[1].ti_le, 1),
            c.tien_do[2].co, c.tien_do[3].co, c.so_cho_chuyen_pdf, c.thieu_thu_muc,
        ])
    ws.append([])
    ws.append([
        "", "TỔNG CỘNG", "", kq.so_dang_vien,
        kq.tien_do[1].co, kq.tien_do[1].tong, round(kq.tien_do[1].ti_le, 1),
        kq.tien_do[2].co, kq.tien_do[3].co, kq.so_cho_chuyen_pdf, kq.thieu_thu_muc,
    ])
    for o in ws[ws.max_row]:
        o.font = Font(bold=True)

    ws2 = wb.create_sheet("Chi_tiet_dang_vien")
    _dat_tieu_de(
        ws2,
        ["ID", "Họ và tên", "Chi bộ", "Tiến độ ƯT1", "Tiến độ ƯT2", "Tiến độ ƯT3",
         "Số tài liệu đạt chuẩn", "Mã tài liệu đã có", "Mã chờ chuyển PDF"],
        [8, 28, 32, 12, 12, 12, 18, 46, 26],
    )
    for d in kq.dong:
        ws2.append([
            d.id, d.ho_ten, d.chi_bo,
            str(d.tien_do[1]), str(d.tien_do[2]), str(d.tien_do[3]),
            d.so_tep_da_co,
            ",".join(str(m) for m in d.da_co),
            ",".join(str(m) for m in d.cho_chuyen_pdf),
        ])

    ws3 = wb.create_sheet("Con_thieu_UT1")
    _dat_tieu_de(
        ws3,
        ["ID", "Họ và tên", "Chi bộ", "Mã tài liệu", "Tên tài liệu còn thiếu"],
        [8, 28, 32, 12, 70],
    )
    dm = danh_muc()
    ma_ut1 = [m.ma for m in sorted(dm.values(), key=lambda x: x.ma) if m.uu_tien == 1]
    for d in kq.dong:
        da_co = set(d.da_co)
        for ma in ma_ut1:
            if ma not in da_co:
                ws3.append([d.id, d.ho_ten, d.chi_bo, ma, dm[ma].ten_day_du])

    thu_muc_ra = Path(thu_muc_ra)
    thu_muc_ra.mkdir(parents=True, exist_ok=True)
    tep = thu_muc_ra / ten_tep_bao_cao(tt, ".xlsx").replace("BaoCao_", "PhuLuc_")
    wb.save(tep)
    return tep
